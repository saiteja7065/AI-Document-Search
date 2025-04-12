from typing import List, Dict, Any
import PyPDF2
from docx import Document
import os
from bs4 import BeautifulSoup
import numpy as np
from config import settings

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = 1000  # characters per chunk
        self.chunk_overlap = 200  # character overlap between chunks
        
    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process a document and return its text content and metadata"""
        text = self._extract_text(file_path, file_type)
        chunks = self._create_chunks(text)
        metadata = self._extract_metadata(file_path, file_type)
        
        return {
            "text": text,
            "chunks": chunks,
            "metadata": metadata
        }
        
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from different file types"""
        if file_type == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_type == "docx":
            return self._extract_from_docx(file_path)
        elif file_type == "txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
        
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
            
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0
        
        while start < len(text):
            # Get chunk with overlap
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Adjust chunk to end at sentence boundary if possible
            if end < len(text):
                last_period = chunk.rfind('.')
                if last_period != -1:
                    chunk = chunk[:last_period + 1]
                    end = start + last_period + 1
            
            chunks.append(chunk.strip())
            
            # Move start position considering overlap
            start = end - self.chunk_overlap
            
            # Ensure we make progress
            if start >= len(text) - self.chunk_overlap:
                break
        
        return chunks
        
    def _extract_metadata(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from the document"""
        file_stats = os.stat(file_path)
        
        metadata = {
            "file_type": file_type,
            "file_size": file_stats.st_size,
            "created_at": file_stats.st_ctime,
            "modified_at": file_stats.st_mtime
        }
        
        # Add file-type specific metadata
        if file_type == "pdf":
            with open(file_path, "rb") as file:
                pdf = PyPDF2.PdfReader(file)
                metadata.update({
                    "page_count": len(pdf.pages),
                    "pdf_info": pdf.metadata if pdf.metadata else {}
                })
        elif file_type == "docx":
            doc = Document(file_path)
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "word_count": sum(len(p.text.split()) for p in doc.paragraphs)
            })
            
        return metadata
        
    def validate_file_type(self, filename: str) -> bool:
        """Check if the file type is supported"""
        extension = filename.lower().split('.')[-1]
        return extension in settings.ALLOWED_FILE_TYPES